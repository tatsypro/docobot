import os
import sys
import re
import time
import requests
import threading
import json
import uuid
from flask import Flask, request as flask_request
from bs4 import BeautifulSoup
from PIL import Image
import pdfplumber
import pytesseract
from docx import Document
from telegram import (
    ReplyKeyboardMarkup,
    ReplyKeyboardRemove,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
)
from telegram.ext import (
    Updater,
    CommandHandler,
    MessageHandler,
    Filters,
    ConversationHandler,
)
from dotenv import load_dotenv
import base64

load_dotenv()

GIGACHAT_AUTH_KEY = "NDEwMjhhYzQtZGRkMS00MDhiLTg3NTktNzE4NmJhOWM4ZWZhOjE3ZTNlOTZjLTE2MzQtNDdiYy04MDUxLWExOWJhYmQ3OTY3Zg=="
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
YANDEX_API_KEY = os.getenv("YANDEX_API_KEY")
OPENAI_KEY = os.getenv("OPENAI_KEY")
IAM_TOKEN = os.getenv("IAM_TOKEN")
FOLDER_ID = os.getenv("FOLDER_ID")
GIGACHAT_CLIENT_ID = os.getenv("GIGACHAT_CLIENT_ID")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET")

print("ID:", GIGACHAT_CLIENT_ID)
print("SECRET:", GIGACHAT_CLIENT_SECRET)
print("AUTH_KEY:", GIGACHAT_AUTH_KEY)

def add_markdown_to_docx(document, text):
    for line in text.split('\n'):
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            document.add_paragraph(line)

# === ЗАМЕНА GPT НА GIGACHAT ===
def query_gigachat(prompt):
    auth_url = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
    chat_url = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"

    headers_auth = {
        "Authorization": f"Basic {GIGACHAT_AUTH_KEY}",
        "RqUID": str(uuid.uuid4()),  # именно UUID, а не os.urandom
        "Content-Type": "application/x-www-form-urlencoded",
        "Accept": "application/json"
    }
    data_auth = {
        "scope": "GIGACHAT_API_PERS"
    }
    print("\n=== ЗАПРОС НА ТОКЕН ===")
    print("auth_url:", auth_url)
    print("headers_auth:", headers_auth)
    print("data_auth:", data_auth)
    print("=======================\n")

    auth_response = requests.post(auth_url, headers=headers_auth, data=data_auth, verify=False)

    print("auth_response.status_code:", auth_response.status_code)
    print("auth_response.text:", auth_response.text)

    access_token = None
    try:
        access_token = auth_response.json().get("access_token")
    except Exception as e:
        print("Ошибка разбора JSON:", e)
        return "❌ Ошибка авторизации GigaChat (токен не получен)"

    if not access_token:
        print("Токен не получен! Проверь параметры запроса.")
        return "❌ Ошибка авторизации GigaChat (токен не получен)"

    headers_chat = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "GigaChat:latest",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "top_p": 0.9,
        "n": 1
    }
    chat_response = requests.post(chat_url, headers=headers_chat, json=payload, verify=False)
    print("chat_response.status_code:", chat_response.status_code)
    print("chat_response.text:", chat_response.text)
    return chat_response.json()['choices'][0]['message']['content']

# 🌐 Flask webhook
app = Flask(__name__)
payment_confirmed_users = set()

@app.route("/webhook", methods=["POST"])
def webhook():
    data = flask_request.form or flask_request.get_json()
    if data.get("Status") == "Completed":
        user_id = data.get("AccountId")
        print(f"✅ Оплата прошла: {user_id}")
        payment_confirmed_users.add(str(user_id))
    return "OK"

@app.route("/")
def index():
    return "Сервер работает ✅"

def run_flask():
    app.run(host="0.0.0.0", port=8080)

threading.Thread(target=run_flask).start()

# 📦 Telegram
LINK, VERIFY, EDIT = range(3)

def start(update, context):
    update.message.reply_text(
        "👋 Привет!\n\n"
        "Мы делаем всё, чтобы оформить комплект документов по 152-ФЗ быстро, удобно и без стресса. "
        "Пожалуйста, пришлите любую из этих вещей — как вам проще:\n\n"
        "1️⃣ Ссылку на ваш сайт (например, tatsy.pro или https://tatsy.pro)\n"
        "2️⃣ Краткие реквизиты:\n"
        "• Наименование юрлица (или ФИО для ИП)\n"
        "• ИНН\n"
        "• ОГРН (или ОГРНИП)\n"
        "• Юридический адрес\n\n"
        "Можете отправить всё одним сообщением, по отдельности или просто приложить файл с реквизитами "
        "(txt, doc, docx, pdf, jpg, png) — бот всё поймёт и аккуратно всё соберёт за вас. "
        "Если что-то забыли — ничего страшного, бот сам подскажет и уточнит, чего не хватает. "
        "Никаких строгих правил и неудобных форм — только забота и поддержка! 🫶"
    )
    context.user_data['pending_fields'] = ['site', 'company', 'inn', 'ogrn', 'address']
    context.user_data['answers'] = {}
    return LINK

def query_yandex(prompt):
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {
        "Authorization": f"Api-Key {IAM_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {
        "modelUri": f"gpt://{FOLDER_ID}/yandexgpt/latest",
        "completionOptions": {
            "stream": False,
            "temperature": 0.3,
            "maxTokens": 2000
        },
        "messages": [{"role": "user", "text": prompt}]
    }
    response = requests.post(url, headers=headers, json=data)
    response.raise_for_status()
    return response.json()['result']['alternatives'][0]['message']['text']

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif ext == ".docx":
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif ext in [".jpg", ".jpeg", ".png"]:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="rus+eng")
        else:
            text = ""
    except Exception as e:
        text = f"[Ошибка при чтении файла: {e}]"
    return text

def handle_link(update, context):
    text = update.message.text.strip()
    lines = text.splitlines()
    site_url = None

    # Найти url сайта
    for line in lines:
        candidate = line.strip()
        if re.match(r'(https?://)?[a-zA-Z0-9\-\_]+\.[a-zA-Z]{2,}', candidate):
            if not candidate.startswith("http"):
                candidate = "https://" + candidate
            site_url = candidate
            break

    # Всё остальное — это реквизиты
    requisites = []
    for l in lines:
        clean_l = l.strip()
        if clean_l and site_url and clean_l != site_url.replace("https://", "") and clean_l != site_url:
            requisites.append(clean_l)
    requisites_text = "\n".join(requisites)
    context.user_data["site_url"] = site_url
    context.user_data["requisites"] = requisites_text

    site_text = ""
    if site_url:
        try:
            r = requests.get(site_url, timeout=10)
            if r.ok:
                soup = BeautifulSoup(r.text, "html.parser")
                site_text = soup.get_text(separator="\n", strip=True)
        except Exception as e:
            site_text = f"[Ошибка при получении сайта: {e}]"

    # --- Используем prompt.txt для сайта, prompt0.txt для всего остального ---
    if site_url:
        with open("prompt.txt", "r", encoding="utf-8") as f:
            base_prompt = f.read()
        full_prompt = (
            base_prompt.strip() +
            f"\n\n🧾 Реквизиты пользователя:\n{requisites_text}\n" +
            f"\n🧾 Текст сайта ({site_url}):\n\n{site_text}"
        )
    else:
        with open("prompt0.txt", "r", encoding="utf-8") as f:
            prompt0 = f.read()
        full_prompt = (
            prompt0.strip() +
            "\n\n---\n"
            "Входные данные пользователя:\n" +
            "\n".join(lines)
        )

    try:
        extracted = query_gigachat(full_prompt)
        update.message.reply_text("🔎 Вот данные, извлечённые из ваших данных:\n\n" + extracted)
        context.user_data["extracted"] = extracted
        keyboard = [["✅ Всё верно", "✏ Изменить"]]
        update.message.reply_text(
            "Проверьте данные выше. Выберите действие:",
            reply_markup=ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
        )
        return VERIFY

    except Exception as e:
        update.message.reply_text(f"❌ Ошибка при анализе данных:\n{e}")
        return LINK

def handle_verification_buttons(update, context):
    text = update.message.text.lower()
    extracted = context.user_data.get("extracted", "")
    if "верно" in text:
        update.message.reply_text(
            "✅ Отлично! Основные реквизиты подтверждены.\n"
            "Сейчас я задам дополнительные вопросы о вашем сайте и услугах.",
            reply_markup=ReplyKeyboardRemove())

        context.user_data["basic_fields_confirmed"] = True
        context.user_data["confirmed_fields"] = extracted

        # Читаем prompt3.txt — здесь будут все уточняющие вопросы!
        with open("prompt3.txt", "r", encoding="utf-8") as f:
            prompt3 = f.read()

        # system_message только из prompt3.txt (без prompt2.txt)
        system_message = {
            "role": "system",
            "content": (
                "Основные реквизиты уже подтверждены:\n" +
                extracted +
                "\n\nНе спрашивай их повторно. Перейди к вопросам ниже по одному:\n\n" +
                prompt3
            )
        }
        context.user_data["chat_history"] = [system_message]

        try:
            prompt = "\n".join(
                [f"{m['role']}: {m['content']}" for m in context.user_data["chat_history"]]
            )
            gpt_reply = query_gigachat(prompt)
            print("[DEBUG] GPT content:", gpt_reply)
            context.user_data["chat_history"].append({
                "role": "assistant",
                "content": gpt_reply
            })
            update.message.reply_text(gpt_reply)
            return EDIT

        except Exception as e:
            update.message.reply_text(f"❌ Ошибка GPT:\n{e}")
            return VERIFY

    elif "изменить" in text:
        context.user_data["awaiting_edit"] = True
        update.message.reply_text(
            "✍️ Что вы хотите изменить? Введите изменения в чат и отправьте.",
            reply_markup=ReplyKeyboardRemove())
        return EDIT

def handle_edit(update, context):
    user_input = update.message.text.strip()
    user_id = str(update.effective_user.id)
    history = context.user_data.get("chat_history", [])
    history.append({"role": "user", "content": user_input})
    with open("prompt2.txt", "r", encoding="utf-8") as file:
        prompt2 = file.read()
    extracted = context.user_data.get("extracted", "")
    system_message = {
        "role": "system",
        "content": f"{prompt2}\n\nВот что уже известно:\n{extracted}"
    }
    try:
        print("\n=== GPT PROMPT (handle_edit, сбор переменных) ===")
        print("system_message:", system_message)
        print("history:", history)
        print("=== КОНЕЦ PROMPT ===\n")
        prompt = "\n".join(
            [f"{m['role']}: {m['content']}" for m in context.user_data["chat_history"]]
        )
        gpt_reply = query_gigachat(prompt)

        # Заменяем маркер [END] на красивую фразу для пользователя
        if "[END]" in gpt_reply:
            gpt_reply = gpt_reply.replace(
                "[END]",
                "Спасибо, все данные собраны! Переходим к оплате и генерации документов."
            )

        # --- ниже по обычной логике
        if "variables" not in context.user_data:
            context.user_data["variables"] = {}

        lines = gpt_reply.split('\n')
        for line in lines:
            match = re.match(r"\[(.+?)\]:\s*(.+)", line)
            if not match:
                match = re.match(r"-?\s*\[(.+?)\]\s*[—-]\s*(.+)", line)
            if match:
                key, value = match.groups()
                context.user_data["variables"][key.strip()] = value.strip()

        history.append({"role": "assistant", "content": gpt_reply})
        context.user_data["chat_history"] = history

        # Теперь проверяем по-прежнему, завершён ли сбор данных (маркер внутри gpt_reply уже заменён)
        if "Спасибо, все данные собраны!" in gpt_reply:
            if user_id in payment_confirmed_users:
                update.message.reply_text(
                    "✅ Отлично! Начинаю генерацию документов...")
                generate_all_documents(update, context)
                return ConversationHandler.END
            else:
                send_payment_button(update, context)
                return EDIT

        update.message.reply_text(gpt_reply)
        return EDIT
    except Exception as e:
        update.message.reply_text(f"❌ Ошибка GPT:\n{e}")
        return EDIT

def send_payment_button(update, context):
    update.message.reply_text(
        "Ура 🥳! Сейчас я сгенерирую документы, и они станут доступны для скачивания в этом чате.\nА пока я работаю, у вас есть время внести оплату 🫶"
    )
    keyboard = [[
        InlineKeyboardButton(
            "💳 Оплатить 4 990 ₽",
            url="https://c.cloudpayments.ru/payments/4abeede5cf154f208314d564f06237d7"
        )
    ]]
    markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text("👇 Нажмите кнопку, чтобы перейти к оплате:",
                              reply_markup=markup)

def generate_all_documents(update, context):
    import time
    from docx import Document
    import os

    print("[DEBUG] Запуск generate_all_documents")
    variables = context.user_data.get("variables", {})
    if variables:
        variables_text = "\n".join([
            f"[{key}]: {value if value else '[НЕ УКАЗАНО]'}"
            for key, value in variables.items()
        ])
    else:
        variables_text = context.user_data.get("extracted", "")
        if not variables_text:
            variables_text = ""

    site_url = context.user_data.get("site_url", "[ваш_сайт]")
    print(f"[DEBUG] site_url: {site_url}")
    print(f"[DEBUG] Текст переменных для промпта:\n{variables_text}")

    documents = [
        {
            "name": "Политика_обработки_ПД.docx",
            "reference_file": "Reference_privacy-policy.txt",
            "title": "🧾 Политика обработки персональных данных"
        },
        {
            "name": "Пользовательское_соглашение.docx",
            "reference_file": "Reference_terms-of-use.txt",
            "title": "📄 Пользовательское соглашение"
        },
        {
            "name": "Согласие_на_обработку_ПД.docx",
            "reference_file": "Reference_сonsent.txt",
            "title": "✍️ Согласие на обработку персональных данных"
        },
        {
            "name": "Публичная_оферта.docx",
            "reference_file": "Reference_offer.txt",
            "title": "💼 Публичная оферта"
        }
    ]

    with open("Prompt_Docobot_API.txt", "r", encoding="utf-8") as f:
        instruction = f.read()
    print("[DEBUG] Прочитан основной промпт из Prompt_Docobot_API.txt")

    for doc in documents:
        try:
            print(f"[DEBUG] Генерируем документ: {doc['name']}")
            with open(doc["reference_file"], "r", encoding="utf-8") as f:
                reference = f.read()
            print(
                f"[DEBUG] Прочитан референс {doc['reference_file']}, длина: {len(reference)} символов"
            )
            prompt = (
                f"{instruction}\n\n"
                f"⬇️ Вот переменные, подтверждённые пользователем:\n"
                f"{variables_text}\n\n"
                f"📎 Вот референс документа:\n"
                f"{reference}\n\n"
                f"{doc['title']}\n"
                "Верни только финальный текст, без вступлений или комментариев."
            )
            print(
                f"[DEBUG] Промпт для GPT (обрезан до 1000 символов):\n{prompt[:1000]}..."
            )
            gpt_reply = query_gigachat(prompt)
            print(f"[DEBUG] Ответ GPT, первые 500 символов:\n{gpt_reply[:500]}...")

            document = Document()
            add_markdown_to_docx(document, gpt_reply)
            document.save(doc["name"])

            print(f"[DEBUG] Документ сохранён: {doc['name']}")
            with open(doc["name"], "rb") as f:
                update.message.reply_document(f)
            os.remove(doc["name"])
            print(f"[DEBUG] Документ удалён после отправки: {doc['name']}")
            time.sleep(1)
        except Exception as e:
            print(f"[ERROR] Ошибка при генерации {doc['name']}: {e}")
            update.message.reply_text(
                f"❌ Ошибка при генерации {doc['name']}:\n{e}")

    update.message.reply_text(
        "📎 Документы сгенерированы!\n\n"
        "✅ Все файлы отправлены в формате .docx\n\n"
        "📌 Что делать дальше:\n"
        "1. Разместите документы на сайте. Обычно ссылки выглядят так:\n"
        f"• Пользовательское соглашение: {site_url}/terms-of-use/\n"
        f"• Политика ПД: {site_url}/privacy-policy/\n"
        f"• Оферта (если есть): {site_url}/offer/\n"
        f"• Согласие на обработку ПД: {site_url}/personal-data-consent/\n\n"
        "2. Добавьте ссылки в подвал сайта.\n"
        "3. При необходимости подайте уведомление в Роскомнадзор:\n"
        "https://pd.rkn.gov.ru/operators-registry/notification/form/\n\n"
        "Если возникнут вопросы — напишите в поддержку: @gestibot_support"
    )

def main():
    updater = Updater(TELEGRAM_TOKEN, use_context=True)
    dp = updater.dispatcher
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            LINK:
            [MessageHandler(Filters.text & ~Filters.command, handle_link)],
            VERIFY: [
                MessageHandler(Filters.text & ~Filters.command,
                               handle_verification_buttons)
            ],
            EDIT:
            [MessageHandler(Filters.text & ~Filters.command, handle_edit)],
        },
        fallbacks=[])
    dp.add_handler(conv)
    updater.start_polling()
    updater.idle()

if __name__ == "__main__":
    main()